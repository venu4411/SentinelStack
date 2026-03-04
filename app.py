"""
SentinelStack v2 - Fixed MongoDB Connection
pip install flask pyjwt python-docx pdfplumber "pymongo[srv]" dnspython
python app.py
"""
import os,json,uuid,hashlib,threading,math,io
from datetime import datetime,timedelta
from functools import wraps
from flask import Flask,request,jsonify,send_file,g
import jwt,urllib.request,ssl

try:
    import docx as docxlib
    from docx.shared import Pt
    HAS_DOCX=True
except:HAS_DOCX=False

BASE_DIR=os.path.dirname(os.path.abspath(__file__))
SECRET_KEY="sentinelstack-v2-jwt"
MONGODB_URI="mongodb+srv://venuwanted36_db_user:Venu%404444@sentinelstack.cdzsevq.mongodb.net/SentinelStack?retryWrites=true&w=majority&appName=SentinelStack"
DB_NAME="SentinelStack"
OPENROUTER_KEY="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"
OPENROUTER_MODEL="openai/gpt-4o"
OPENROUTER_URL="https://openrouter.ai/api/v1/chat/completions"
SSL_CTX=ssl.create_default_context()
app=Flask(__name__)
app.config['SECRET_KEY']=SECRET_KEY

# ── MongoDB ───────────────────────────────────────────────────
_db=None
USE_MONGO=False

def connect_mongo():
    global _db,USE_MONGO
    try:
        from pymongo import MongoClient
        from pymongo.server_api import ServerApi
        print("  Connecting to MongoDB Atlas...")
        client=MongoClient(
            MONGODB_URI,
            server_api=ServerApi('1'),
            serverSelectionTimeoutMS=15000,
            connectTimeoutMS=15000,
            tls=True
        )
        client.admin.command('ping')
        _db=client[DB_NAME]
        USE_MONGO=True
        print(f"  MongoDB Atlas CONNECTED — database: {DB_NAME}")
        return True
    except ImportError:
        print("  pymongo not installed. Run: pip install 'pymongo[srv]' dnspython")
    except Exception as e:
        print(f"  MongoDB error: {e}")
    USE_MONGO=False
    print("  Falling back to local SQLite")
    return False

def col(name):
    return _db[name] if USE_MONGO and _db is not None else None

def db_find_one(c,q):
    cl=col(c)
    if cl is not None:
        d=cl.find_one(q)
        if d:d['id']=str(d.pop('_id',d.get('id','')))
        return d
    return _sq_find_one(c,q)

def db_find(c,q=None,sf=None,sd=-1):
    cl=col(c)
    if cl is not None:
        cur=cl.find(q or {})
        if sf:cur=cur.sort(sf,sd)
        r=[]
        for d in cur:
            d['id']=str(d.pop('_id',d.get('id','')))
            r.append(d)
        return r
    return _sq_find(c,q)

def db_insert(c,doc):
    if '_id' not in doc:doc['_id']=doc.get('id',str(uuid.uuid4()))
    cl=col(c)
    if cl is not None:
        try:cl.insert_one(dict(doc))
        except Exception as e:
            if 'duplicate' not in str(e).lower():raise
    else:_sq_insert(c,doc)
    doc['id']=str(doc.pop('_id',doc.get('id','')))
    return doc['id']

def db_update(c,q,u):
    cl=col(c)
    if cl is not None:cl.update_one(q,{'$set':u})
    else:_sq_update(c,q,u)

def db_count(c,q=None):
    cl=col(c)
    if cl is not None:return cl.count_documents(q or {})
    return len(_sq_find(c,q))

# ── SQLite fallback ───────────────────────────────────────────
import sqlite3
_sq=None
SQ_PATH=os.path.join(BASE_DIR,'sentinel_local.db')
COLS=['users','questionnaires','questions','reference_documents','document_chunks','runs','answers','answer_versions','audit_logs']

def get_sq():
    global _sq
    if _sq is None:
        _sq=sqlite3.connect(SQ_PATH,check_same_thread=False)
        _sq.row_factory=sqlite3.Row
        _sq.execute("PRAGMA journal_mode=WAL")
        for t in COLS:_sq.execute(f"CREATE TABLE IF NOT EXISTS {t}(id TEXT PRIMARY KEY,data TEXT)")
        _sq.commit()
    return _sq

def _sq_find_one(c,q):
    try:
        rows=get_sq().execute(f"SELECT data FROM {c}").fetchall()
        for r in rows:
            d=json.loads(r['data'])
            if not q or all(d.get(k)==v for k,v in q.items()):return d
    except:pass
    return None

def _sq_find(c,q=None):
    try:
        rows=get_sq().execute(f"SELECT data FROM {c}").fetchall()
        result=[]
        for r in rows:
            d=json.loads(r['data'])
            if not q or all(d.get(k)==v for k,v in q.items()):result.append(d)
        return result
    except:return []

def _sq_insert(c,doc):
    db=get_sq()
    rid=str(doc.get('_id') or doc.get('id') or uuid.uuid4())
    doc['id']=rid;doc.pop('_id',None)
    db.execute(f"INSERT OR REPLACE INTO {c}(id,data) VALUES(?,?)",(rid,json.dumps(doc)))
    db.commit()

def _sq_update(c,q,u):
    db=get_sq()
    rows=db.execute(f"SELECT rowid,data FROM {c}").fetchall()
    for r in rows:
        d=json.loads(r['data'])
        if all(d.get(k)==v for k,v in q.items()):
            d.update(u)
            db.execute(f"UPDATE {c} SET data=? WHERE rowid=?",(json.dumps(d),r['rowid']))
    db.commit()

# ── Reference docs ────────────────────────────────────────────
REFERENCE_DOCS={
"Security Policy":"""SentinelStack Security Policy v2.1
DATA ENCRYPTION: AES-256 at rest. TLS 1.3 in transit. AWS KMS quarterly key rotation. PII field-level encryption.
ACCESS CONTROL: Least privilege. MFA mandatory for all production access. PAM deployed. Quarterly access reviews.
VULNERABILITY: Monthly scans. Critical/high patched within 30 days. Annual pen test. Bug bounty active.
MONITORING: 24/7 SOC. SIEM 90-day retention. Alerts within 15 minutes business hours.
TRAINING: Annual security awareness. Quarterly phishing tests. Background checks. Pre-access training for new hires.""",

"Data Retention Policy":"""SentinelStack Data Retention Policy v1.4
CUSTOMER DATA: Contract duration plus 90 days post-termination. DoD 5220.22-M deletion. 30-day self-service deletion requests.
BACKUPS: 30-day retention encrypted cloud. Point-in-time recovery 7 days. Quarterly restoration tests.
LOGS: Application 90 days. Security audit 2 years. Access 1 year. Compliance 7 years.
GDPR: Access, rectification, erasure, portability. Acknowledged 72 hours, fulfilled 30 days. DPO: privacy@sentinelstack.com.""",

"Incident Response Plan":"""SentinelStack Incident Response Plan v3.0
P0 Critical: breach ransomware complete outage. Response 15 minutes.
P1 High: significant degradation suspected breach. Response 1 hour.
P2 Medium: partial impact. Response 4 hours. P3 Low: 24 hours.
PHASES: Detection 0-15 min. Containment 15 min-2 hr. Eradication 2-24 hr. Recovery 24-72 hr. Post-incident 2 weeks.
NOTIFICATION: Regulatory GDPR 72 hours confirmed breach. Customer 48 hours email and dashboard.""",

"Compliance Certifications":"""SentinelStack Compliance v2.0
SOC 2 TYPE II: Security Availability Processing Integrity Confidentiality Privacy. Deloitte Touche. Oct 2022-Sep 2023.
ISO 27001:2022: BSI Group certified 2021. Certificate IS 754821.
GDPR: DPAs for EU customers. EU-US Data Privacy Framework certified.
HIPAA: BAAs available. PHI not stored.
PCI DSS: SAQ-D. Annual QSA. Cardholder data never stored.
PENTEST: CrowdStrike November 2023. No critical findings. FedRAMP Moderate Q4 2024.""",

"Access Control Policy":"""SentinelStack Access Control Policy v2.3
IAM: Okta SSO. RBAC all systems. No shared production accounts.
AUTHENTICATION: 14-character minimum. MFA required production portals VPN. TOTP FIDO2 supported. 8-hour timeout 2-hour admin.
PAM: CyberArk. JIT provisioning. Sessions recorded 1 year. Monthly reviews.
PROVISIONING: New employee 1 business day. Termination 2 hours. Contractor end date plus 1 day. Inactive 90 days suspended.
THIRD PARTY: Per-project signed DPA. Dedicated VPN MFA. Monthly log review."""
}

SAMPLE_Q=[
    {"order":1,"text":"Does your organization maintain any security certifications?","category":"Certifications"},
    {"order":2,"text":"How is data encrypted at rest and in transit?","category":"Data Security"},
    {"order":3,"text":"What is your incident response time for critical security incidents?","category":"Incident Response"},
    {"order":4,"text":"Do you support multi-factor authentication (MFA)?","category":"Access Control"},
    {"order":5,"text":"What is your data retention policy after contract termination?","category":"Data Retention"},
    {"order":6,"text":"How do you handle GDPR data subject access requests?","category":"Privacy"},
    {"order":7,"text":"Describe your vulnerability management and patching process.","category":"Vulnerability"},
    {"order":8,"text":"What are your breach notification timelines?","category":"Incident Response"},
    {"order":9,"text":"How is privileged access managed?","category":"Access Control"},
    {"order":10,"text":"Do you conduct employee security awareness training?","category":"Training"},
    {"order":11,"text":"What is your backup and disaster recovery process?","category":"Business Continuity"},
    {"order":12,"text":"How do you manage third-party vendor access?","category":"Third-Party Risk"},
]

# ── RAG ───────────────────────────────────────────────────────
def chunk_text(text,size=200,overlap=25):
    words=text.split();chunks=[];i=0
    while i<len(words):
        chunks.append(" ".join(words[i:i+size]));i+=size-overlap
    return[c for c in chunks if len(c)>15]

def embed(t):
    t=t.lower();v=[0.0]*256
    for i in range(len(t)-2):v[hash(t[i:i+3])%256]+=1.0
    n=math.sqrt(sum(x*x for x in v)) or 1e-9
    return[x/n for x in v]

def cosine(a,b):
    return sum(x*y for x,y in zip(a,b))/(math.sqrt(sum(x*x for x in a)) or 1e-9)/(math.sqrt(sum(x*x for x in b)) or 1e-9)

def bm25(qt,dt,adl,k1=1.5,b=0.75):
    from collections import Counter
    tf=Counter(dt);dl=len(dt);s=0.0
    for t in qt:
        f=tf.get(t,0)
        if f:s+=math.log(2)*(f*(k1+1))/(f+k1*(1-b+b*dl/max(adl,1)))
    return s

def hybrid_search(q,k=5):
    chunks=[]
    for dn,dt in REFERENCE_DOCS.items():
        for i,c in enumerate(chunk_text(dt)):chunks.append({'text':c,'document':dn,'idx':i})
    adl=sum(len(c['text'].split()) for c in chunks)/max(len(chunks),1)
    qv=embed(q);qt=q.lower().split()
    for c in chunks:
        c['similarity']=round(0.65*cosine(qv,embed(c['text']))+0.35*min(bm25(qt,c['text'].lower().split(),adl)/8,1.0),4)
    chunks.sort(key=lambda x:x['similarity'],reverse=True)
    return chunks[:k]

SYSP="""You are a compliance expert for SentinelStack. Answer questions using ONLY the provided reference excerpts.
RULES: 1) Only use info from references 2) Always cite document names 3) If insufficient, return exactly: "Not found in references." 4) No hallucination.
Return ONLY valid JSON: {"answer":"...","citations":["DocName"],"evidence_snippets":["excerpt"],"confidence_score":0.0}"""

def call_ai(messages):
    data=json.dumps({"model":OPENROUTER_MODEL,"messages":messages,"max_tokens":700,"temperature":0.1,"response_format":{"type":"json_object"}}).encode()
    req=urllib.request.Request(OPENROUTER_URL,data=data,headers={"Authorization":f"Bearer {OPENROUTER_KEY}","Content-Type":"application/json","HTTP-Referer":"https://sentinelstack.ai"})
    r=urllib.request.urlopen(req,context=SSL_CTX,timeout=30)
    return json.loads(r.read())['choices'][0]['message']['content']

def conf(sims,cn):
    if not sims:return 0.0,"None"
    avg=sum(sims)/len(sims);div=len(set(round(s,1) for s in sims))/max(len(sims),1);cf=min(cn/3,1.0)
    sc=round(avg*0.5+div*0.2+cf*0.3,2)
    return sc,("High" if sc>=0.72 else "Medium" if sc>=0.52 else "Low" if sc>=0.35 else "None")

def rag(question):
    chunks=hybrid_search(question)
    best=chunks[0]['similarity'] if chunks else 0
    if best<0.42:
        return{"answer":"Not found in references.","citations":[],"evidence_snippets":[],"confidence_score":0.0,"confidence_label":"None","status":"not_found","chunks":chunks}
    ctx="\n\n---\n\n".join(f"[Source:{c['document']}]\n{c['text']}" for c in chunks)
    try:
        raw=call_ai([{"role":"system","content":SYSP},{"role":"user","content":f"References:\n{ctx}\n\nQuestion: {question}\n\nReturn JSON only."}])
        res=json.loads(raw.strip().lstrip('```json').rstrip('```').strip())
        ans=res.get("answer","Not found in references.")
        cits=res.get("citations",[])
        snips=res.get("evidence_snippets",[])
        if not cits:ans="Not found in references."
        sims=[c['similarity'] for c in chunks[:max(len(cits),1)]]
        sc,lb=conf(sims,len(cits))
        return{"answer":ans,"citations":cits,"evidence_snippets":snips[:3],"confidence_score":sc,"confidence_label":lb,"status":"generated" if ans!="Not found in references." else "not_found","chunks":chunks}
    except Exception as e:
        print(f"  AI error: {e}"); b=chunks[0]
        sims=[c['similarity'] for c in chunks[:2]];sc,lb=conf(sims,1)
        return{"answer":f"Based on {b['document']}: {b['text'][:300]}","citations":[b['document']],"evidence_snippets":[b['text'][:200]],"confidence_score":sc,"confidence_label":lb,"status":"generated","chunks":chunks}

# ── Helpers ───────────────────────────────────────────────────
def phash(p):return hashlib.sha256((p+"_sentinel_v2").encode()).hexdigest()
def now():return datetime.utcnow().isoformat()
def nid():return str(uuid.uuid4())

def make_token(uid,email,role):
    return jwt.encode({'sub':uid,'email':email,'role':role,'exp':datetime.utcnow()+timedelta(hours=24)},SECRET_KEY,algorithm='HS256')

def require_auth(f):
    @wraps(f)
    def inner(*a,**kw):
        auth=request.headers.get('Authorization','')
        if not auth.startswith('Bearer '):return jsonify({'error':'Unauthorized'}),401
        try:
            p=jwt.decode(auth.split()[1],SECRET_KEY,algorithms=['HS256'])
            g.uid=p['sub'];g.email=p.get('email','');g.role=p.get('role','user')
        except:return jsonify({'error':'Invalid token'}),401
        return f(*a,**kw)
    return inner

def audit(action,rtype=None,rid=None,meta=None):
    try:db_insert('audit_logs',{'_id':nid(),'user_id':getattr(g,'uid',None),'user_email':getattr(g,'email',''),'action':action,'resource_type':rtype,'resource_id':rid,'metadata':meta or {},'created_at':now()})
    except:pass

def ser(d):
    if not d:return None
    d=dict(d)
    if '_id' in d:d['id']=str(d.pop('_id'))
    return d

# ── CORS ──────────────────────────────────────────────────────
@app.after_request
def cors(r):
    r.headers['Access-Control-Allow-Origin']='*'
    r.headers['Access-Control-Allow-Headers']='Content-Type,Authorization'
    r.headers['Access-Control-Allow-Methods']='GET,POST,PUT,PATCH,DELETE,OPTIONS'
    return r

@app.before_request
def preflight():
    if request.method=='OPTIONS':
        from flask import Response;return Response(status=200)

# ── Seed ──────────────────────────────────────────────────────
_seeded=False
def seed():
    global _seeded
    if _seeded:return
    _seeded=True
    if db_find_one('users',{'email':'admin@sentinelstack.com'}):
        print("  Database already seeded ✓");return
    print("  Seeding database...")
    aid=nid();did_=nid()
    db_insert('users',{'_id':aid,'email':'admin@sentinelstack.com','name':'Admin User','password_hash':phash('admin123!'),'role':'admin','created_at':now()})
    db_insert('users',{'_id':did_,'email':'demo@sentinelstack.com','name':'Demo User','password_hash':phash('demo123!'),'role':'user','created_at':now()})
    for dn,dt in REFERENCE_DOCS.items():
        rid=nid();cks=chunk_text(dt)
        db_insert('reference_documents',{'_id':rid,'owner_id':aid,'name':dn,'description':f'SentinelStack {dn}','doc_type':'policy','chunk_count':len(cks),'is_indexed':True,'created_at':now()})
        for i,c in enumerate(cks):db_insert('document_chunks',{'_id':nid(),'document_id':rid,'document_name':dn,'chunk_index':i,'text':c})
    qid=nid()
    db_insert('questionnaires',{'_id':qid,'owner_id':did_,'title':'Enterprise Security Assessment 2024','description':'Annual vendor security questionnaire','status':'ready','created_at':now()})
    for q in SAMPLE_Q:db_insert('questions',{'_id':nid(),'questionnaire_id':qid,'order_index':q['order'],'text':q['text'],'category':q['category']})
    print(f"  Seeded: users, {len(REFERENCE_DOCS)} docs, questionnaire, {len(SAMPLE_Q)} questions ✓")

@app.before_request
def auto_seed():seed()

# ── Auth ──────────────────────────────────────────────────────
@app.route('/api/auth/login',methods=['POST'])
def login():
    d=request.json or {}
    user=db_find_one('users',{'email':d.get('email','')})
    if not user or user.get('password_hash')!=phash(d.get('password','')):
        return jsonify({'error':'Invalid credentials'}),401
    uid=str(user.get('id') or user.get('_id',''))
    token=make_token(uid,user['email'],user['role'])
    g.uid=uid;g.email=user['email']
    audit('LOGIN','auth',uid)
    return jsonify({'access_token':token,'token_type':'bearer','user':{'id':uid,'email':user['email'],'name':user['name'],'role':user['role']}})

@app.route('/api/auth/register',methods=['POST'])
def register():
    d=request.json or {}
    if not all([d.get('email'),d.get('password'),d.get('name')]):return jsonify({'error':'All fields required'}),400
    if db_find_one('users',{'email':d['email']}):return jsonify({'error':'Email already registered'}),400
    uid=nid()
    db_insert('users',{'_id':uid,'email':d['email'],'name':d['name'],'password_hash':phash(d['password']),'role':'user','created_at':now()})
    return jsonify({'access_token':make_token(uid,d['email'],'user'),'token_type':'bearer','user':{'id':uid,'email':d['email'],'name':d['name'],'role':'user'}})

@app.route('/api/auth/me')
@require_auth
def me():
    u=db_find_one('users',{'id':g.uid}) or db_find_one('users',{'_id':g.uid})
    return jsonify(u or {})

# ── Questionnaires ────────────────────────────────────────────
@app.route('/api/questionnaires')
@require_auth
def list_qs():
    qs=db_find('questionnaires',{'owner_id':g.uid})
    result=[]
    for q in qs:
        qid=str(q.get('id') or q.get('_id',''))
        qc=db_count('questions',{'questionnaire_id':qid})
        runs=db_find('runs',{'questionnaire_id':qid},sf='version')
        lat=runs[-1] if runs else None
        pct=round(lat.get('answered',0)/max(lat.get('total',1),1)*100,1) if lat else 0
        result.append({**ser(q),'question_count':qc,'latest_run':lat.get('version') if lat else None,'completion_pct':pct})
    return jsonify(result)

@app.route('/api/questionnaires/<qid>')
@require_auth
def get_q(qid):
    q=db_find_one('questionnaires',{'id':qid}) or db_find_one('questionnaires',{'_id':qid})
    if not q:return jsonify({'error':'Not found'}),404
    qs=sorted(db_find('questions',{'questionnaire_id':qid}),key=lambda x:x.get('order_index',0))
    runs=sorted(db_find('runs',{'questionnaire_id':qid}),key=lambda x:x.get('version',0),reverse=True)
    d=ser(q);d['questions']=[ser(x) for x in qs];d['runs']=[ser(x) for x in runs]
    return jsonify(d)

@app.route('/api/questionnaires',methods=['POST'])
@require_auth
def create_q():
    d=request.json or {}
    if not d.get('title'):return jsonify({'error':'Title required'}),400
    qid=nid()
    db_insert('questionnaires',{'_id':qid,'owner_id':g.uid,'title':d['title'],'description':d.get('description',''),'status':'ready','created_at':now()})
    for i,item in enumerate(d.get('questions',[])):
        db_insert('questions',{'_id':nid(),'questionnaire_id':qid,'order_index':i+1,'text':item['text'],'category':item.get('category','General')})
    audit('CREATE_QUESTIONNAIRE','questionnaire',qid)
    return jsonify({'id':qid,'title':d['title']})

# ── Runs ──────────────────────────────────────────────────────
def bg_run(qid,run_id,q_ids=None):
    try:
        qs=sorted(db_find('questions',{'questionnaire_id':qid}),key=lambda x:x.get('order_index',0))
        if q_ids:qs=[q for q in qs if str(q.get('id') or q.get('_id','')) in q_ids]
        ans=nf=hc=lc=0;sims=[]
        print(f"\n  GPT-4o processing {len(qs)} questions...")
        for q in qs:
            q_id=str(q.get('id') or q.get('_id',''))
            print(f"  Q{q.get('order_index')}: {q.get('text','')[:50]}...")
            res=rag(q.get('text',''))
            ex=db_find_one('answers',{'question_id':q_id,'run_id':run_id})
            if ex:
                aid=str(ex.get('id') or ex.get('_id',''))
                vn=db_count('answer_versions',{'answer_id':aid})
                db_insert('answer_versions',{'_id':nid(),'answer_id':aid,'answer_text':ex.get('answer_text',''),'citations':ex.get('citations',[]),'confidence_score':ex.get('confidence_score',0),'version_number':vn+1,'change_reason':'regenerated','created_at':now()})
                db_update('answers',{'id':aid},{'answer_text':res['answer'],'citations':res['citations'],'evidence_snippets':res['evidence_snippets'],'confidence_score':res['confidence_score'],'confidence_label':res['confidence_label'],'status':res['status'],'similarity_scores':[c['similarity'] for c in res.get('chunks',[])],'updated_at':now()})
            else:
                db_insert('answers',{'_id':nid(),'question_id':q_id,'run_id':run_id,'answer_text':res['answer'],'citations':res['citations'],'evidence_snippets':res['evidence_snippets'],'confidence_score':res['confidence_score'],'confidence_label':res['confidence_label'],'status':res['status'],'similarity_scores':[c['similarity'] for c in res.get('chunks',[])],'is_manually_edited':False,'edit_note':'','created_at':now(),'updated_at':now()})
            if res['status']=='generated':
                ans+=1
                if res['confidence_label']=='High':hc+=1
                elif res['confidence_label'] in('Low','None'):lc+=1
            else:nf+=1
            sims.extend([c['similarity'] for c in res.get('chunks',[])])
        avg=round(sum(sims)/len(sims),3) if sims else 0.0
        db_update('runs',{'id':run_id},{'status':'completed','total':len(qs),'answered':ans,'not_found':nf,'high_confidence':hc,'low_confidence':lc,'avg_similarity':avg,'completed_at':now()})
        print(f"  Run complete: {ans}/{len(qs)} | avg_sim={avg}")
    except Exception as e:
        db_update('runs',{'id':run_id},{'status':'failed','error':str(e)})
        import traceback;traceback.print_exc()

@app.route('/api/questionnaires/<qid>/run',methods=['POST'])
@require_auth
def start_run(qid):
    q=db_find_one('questionnaires',{'id':qid}) or db_find_one('questionnaires',{'_id':qid})
    if not q:return jsonify({'error':'Not found'}),404
    v=max((r.get('version',0) for r in db_find('runs',{'questionnaire_id':qid})),default=0)+1
    run_id=nid()
    db_insert('runs',{'_id':run_id,'questionnaire_id':qid,'version':v,'status':'running','model_used':OPENROUTER_MODEL,'total':0,'answered':0,'not_found':0,'high_confidence':0,'low_confidence':0,'avg_similarity':0.0,'created_at':now()})
    audit('START_RUN','run',run_id,{'version':v})
    d=request.json or {}
    threading.Thread(target=bg_run,args=(qid,run_id,d.get('question_ids')),daemon=True).start()
    return jsonify({'run_id':run_id,'version':v,'status':'running'})

@app.route('/api/runs/<run_id>')
@require_auth
def get_run(run_id):
    run=db_find_one('runs',{'id':run_id}) or db_find_one('runs',{'_id':run_id})
    if not run:return jsonify({'error':'Not found'}),404
    ans_raw=db_find('answers',{'run_id':run_id})
    answers=[]
    for a in ans_raw:
        q_id=a.get('question_id','')
        q=db_find_one('questions',{'id':q_id}) or db_find_one('questions',{'_id':q_id})
        d=ser(a);d['question_text']=q.get('text','') if q else '';d['category']=q.get('category','') if q else '';d['order']=q.get('order_index',0) if q else 0
        answers.append(d)
    answers.sort(key=lambda x:x.get('order',0))
    r=ser(run);r['answers']=answers
    return jsonify(r)

@app.route('/api/answers/<aid>',methods=['PATCH'])
@require_auth
def update_ans(aid):
    old=db_find_one('answers',{'id':aid}) or db_find_one('answers',{'_id':aid})
    if not old:return jsonify({'error':'Not found'}),404
    d=request.json or {}
    vn=db_count('answer_versions',{'answer_id':aid})
    db_insert('answer_versions',{'_id':nid(),'answer_id':aid,'answer_text':old.get('answer_text',''),'citations':old.get('citations',[]),'confidence_score':old.get('confidence_score',0),'version_number':vn+1,'change_reason':d.get('edit_note','manual edit'),'created_at':now()})
    u={'updated_at':now()}
    if 'answer_text' in d:u.update({'answer_text':d['answer_text'],'is_manually_edited':True,'edit_note':d.get('edit_note','')})
    if 'citations' in d:u['citations']=d['citations']
    if 'status' in d:u['status']=d['status']
    db_update('answers',{'id':aid},u)
    audit('EDIT_ANSWER','answer',aid)
    return jsonify({'success':True})

@app.route('/api/answers/<aid>/versions')
@require_auth
def ans_versions(aid):
    vers=sorted(db_find('answer_versions',{'answer_id':aid}),key=lambda x:x.get('version_number',0),reverse=True)
    return jsonify([ser(v) for v in vers])

# ── Documents ─────────────────────────────────────────────────
@app.route('/api/documents')
@require_auth
def list_docs():
    return jsonify([ser(d) for d in db_find('reference_documents',{})])

@app.route('/api/documents/upload',methods=['POST'])
@require_auth
def upload_doc():
    name=request.form.get('name','').strip()
    desc=request.form.get('description','').strip()
    file=request.files.get('file')
    if not file or not name:return jsonify({'error':'file and name required'}),400
    content=file.read();fname=file.filename.lower()
    try:
        if fname.endswith('.pdf'):
            import pdfplumber,io as _io
            with pdfplumber.open(_io.BytesIO(content)) as pdf:text="\n".join(p.extract_text() or'' for p in pdf.pages)
        elif fname.endswith('.docx'):
            import docx as _d,io as _io;text="\n".join(p.text for p in _d.Document(_io.BytesIO(content)).paragraphs if p.text.strip())
        else:text=content.decode('utf-8',errors='ignore')
    except Exception as e:return jsonify({'error':f'Parse error:{e}'}),400
    if not text.strip():return jsonify({'error':'No text'}),400
    did=nid();cks=chunk_text(text)
    REFERENCE_DOCS[name]=text
    db_insert('reference_documents',{'_id':did,'owner_id':g.uid,'name':name,'description':desc,'doc_type':fname.rsplit('.',1)[-1],'chunk_count':len(cks),'is_indexed':True,'created_at':now()})
    for i,c in enumerate(cks):db_insert('document_chunks',{'_id':nid(),'document_id':did,'document_name':name,'chunk_index':i,'text':c})
    audit('UPLOAD_DOCUMENT','document',did,{'name':name,'chunks':len(cks)})
    return jsonify({'id':did,'name':name,'chunk_count':len(cks)})

# ── Export ────────────────────────────────────────────────────
@app.route('/api/runs/<run_id>/export/docx')
@require_auth
def export_docx(run_id):
    if not HAS_DOCX:return jsonify({'error':'python-docx not installed'}),500
    run=db_find_one('runs',{'id':run_id}) or db_find_one('runs',{'_id':run_id})
    if not run:return jsonify({'error':'Not found'}),404
    qobj=db_find_one('questionnaires',{'id':run.get('questionnaire_id','')})
    answers=db_find('answers',{'run_id':run_id})
    doc=docxlib.Document()
    doc.add_heading('SentinelStack AI Compliance Response',0)
    doc.add_paragraph(f"Questionnaire: {qobj.get('title','') if qobj else ''}")
    doc.add_paragraph(f"Run v{run.get('version',1)} — {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph(f"AI Model: {run.get('model_used',OPENROUTER_MODEL)}")
    doc.add_page_break();doc.add_heading('Responses',level=1)
    for a in sorted(answers,key=lambda x:x.get('order',0)):
        q_id=a.get('question_id','')
        q=db_find_one('questions',{'id':q_id}) or db_find_one('questions',{'_id':q_id})
        cits=a.get('citations',[]);snips=a.get('evidence_snippets',[])
        doc.add_heading(f"Q{q.get('order_index','?') if q else '?'}. {q.get('text','') if q else ''}",level=2)
        doc.add_paragraph(f"Confidence: {a.get('confidence_label','N/A')} ({round(a.get('confidence_score',0)*100)}%)")
        p=doc.add_paragraph();p.add_run("Answer: ").bold=True;p.add_run(a.get('answer_text',''))
        if cits:p2=doc.add_paragraph();p2.add_run("Citations: ").bold=True;p2.add_run(" | ".join(cits))
        if snips:p3=doc.add_paragraph();p3.add_run("Evidence: ").bold=True;p3.add_run(snips[0][:200])
        doc.add_paragraph("─"*50)
    buf=io.BytesIO();doc.save(buf);buf.seek(0)
    audit('EXPORT_DOCX','run',run_id)
    return send_file(buf,mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',as_attachment=True,download_name=f'sentinelstack_v{run.get("version",1)}.docx')

# ── Analytics + Audit ─────────────────────────────────────────
@app.route('/api/analytics/overview')
@require_auth
def analytics():
    tq=db_count('questionnaires',{'owner_id':g.uid});td=db_count('reference_documents',{})
    aa=db_find('answers',{});ta=len(aa);hc=len([a for a in aa if a.get('confidence_label')=='High'])
    runs=db_find('runs',{});avg=round(sum(r.get('avg_similarity',0) for r in runs)/max(len(runs),1),3) if runs else 0
    return jsonify({'total_questionnaires':tq,'total_documents':td,'total_answers':ta,'high_confidence_rate':round(hc/max(ta,1)*100,1),'avg_similarity':avg,'total_runs':len(runs),'database':'MongoDB Atlas' if USE_MONGO else 'SQLite fallback'})

@app.route('/api/audit-logs')
@require_auth
def audit_logs():
    logs=db_find('audit_logs',None if g.role=='admin' else {'user_id':g.uid},sf='created_at')
    return jsonify([ser(l) for l in reversed(logs[-100:])])

@app.route('/api/health')
def health():
    return jsonify({'status':'ok','service':'SentinelStack v2','database':'MongoDB Atlas' if USE_MONGO else 'SQLite fallback','ai_model':OPENROUTER_MODEL,'time':now()})

@app.route('/')
def index():
    return send_file(os.path.join(BASE_DIR,'index.html'))

if __name__=='__main__':
    print("━"*55)
    print("  SentinelStack v2 — AI Compliance Platform")
    print("━"*55)
    connect_mongo()
    seed()
    print(f"\n  Database : {'MongoDB Atlas (cloud persistent) ✅' if USE_MONGO else 'SQLite local fallback ⚠'}")
    print(f"  AI Model : {OPENROUTER_MODEL} via OpenRouter")
    print(f"  URL      : http://localhost:7860")
    print(f"  Login    : demo@sentinelstack.com / demo123!")
    print("━"*55)
    app.run(host='0.0.0.0',port=7860,debug=False,threaded=True)